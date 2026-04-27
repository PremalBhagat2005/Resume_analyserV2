import io
import os
import re
from PyPDF2 import PdfReader
from docx import Document
from werkzeug.utils import secure_filename
from app.utils.helpers import clean_text


def parse_pdf(file):
    """
    Extract text from PDF file.
    Returns extracted text string.
    """
    file.seek(0)
    pdf_bytes = file.read()

    try:
        import pdfplumber

        text_parts = []

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                words = page.extract_words(
                    x_tolerance=3,
                    y_tolerance=3,
                    keep_blank_chars=False,
                    use_text_flow=True,
                    horizontal_ltr=True,
                    vertical_ttb=True,
                    extra_attrs=["size", "fontname"],
                )

                if not words:
                    page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                    if page_text:
                        text_parts.append(page_text)
                    continue

                # Group words by nearby y-coordinate to reconstruct readable lines.
                lines = {}
                for word in words:
                    y_key = round(word["top"] / 3) * 3
                    lines.setdefault(y_key, []).append(word)

                for y_key in sorted(lines.keys()):
                    line_words = sorted(lines[y_key], key=lambda w: w["x0"])
                    line_text = " ".join(w["text"] for w in line_words)
                    text_parts.append(line_text)

        full_text = "\n".join(text_parts)
        if not full_text:
            return ""

        cleaned_lines = []
        for line in full_text.splitlines():
            line = re.sub(r'[^\w\s.,;:\'"@()\-#]', '', line).strip()
            line = re.sub(r'\s+', ' ', line).strip()
            if line:
                cleaned_lines.append(line)

        return "\n".join(cleaned_lines)

    except Exception as e:
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract_text

            text = pdfminer_extract_text(io.BytesIO(pdf_bytes))
            if not text:
                return ""

            cleaned_lines = []
            for line in text.splitlines():
                line = re.sub(r'[^\w\s.,;:\'"@()\-#]', '', line).strip()
                line = re.sub(r'\s+', ' ', line).strip()
                if line:
                    cleaned_lines.append(line)

            return "\n".join(cleaned_lines)
        except Exception as e2:
            try:
                # Last-resort fallback
                pdf_reader = PdfReader(io.BytesIO(pdf_bytes))
                text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                if not text:
                    return ""

                cleaned_lines = []
                for line in text.splitlines():
                    line = re.sub(r'[^\w\s.,;:\'"@()\-#]', '', line).strip()
                    line = re.sub(r'\s+', ' ', line).strip()
                    if line:
                        cleaned_lines.append(line)

                return "\n".join(cleaned_lines)
            except Exception as e3:
                raise ValueError(
                    f"Could not parse PDF: {str(e)} | Fallback error: {str(e2)} | PyPDF2 error: {str(e3)}"
                )


def parse_docx(file):
    """
    Extract text from DOCX file.
    Returns extracted text string.
    """
    try:
        file.seek(0)
        doc = Document(io.BytesIO(file.read()))

        sections_text = []
        seen_lines = set()
        largest_font_size = 0
        name_candidate = None

        # --- PASS 1: Read paragraphs with font size awareness ---
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            max_size_in_para = 0
            for run in para.runs:
                if run.font.size:
                    size_pt = run.font.size.pt
                    if size_pt > max_size_in_para:
                        max_size_in_para = size_pt

            if max_size_in_para > largest_font_size and max_size_in_para > 14:
                words = text.split()
                if 2 <= len(words) <= 5 and not any(c in text for c in ['@', '/', '+']):
                    largest_font_size = max_size_in_para
                    name_candidate = text

            if text not in seen_lines:
                sections_text.append(text)
                seen_lines.add(text)

        # --- PASS 2: Extract text from tables (common for column-layout resumes) ---
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        line = para.text.strip()
                        if line and line not in seen_lines:
                            sections_text.append(line)
                            seen_lines.add(line)

        # --- PASS 3: Extract text box content from document XML ---
        try:
            for node in doc.element.body.iter():
                if node.tag.endswith('}txbxContent'):
                    for tnode in node.iter():
                        if tnode.tag.endswith('}t') and tnode.text:
                            line = tnode.text.strip()
                            if line and line not in seen_lines:
                                sections_text.append(line)
                                seen_lines.add(line)
        except Exception:
            pass

        # --- PASS 4: Extract from headers and footers ---
        try:
            for section in doc.sections:
                header_blocks = [section.header, section.first_page_header, section.even_page_header]
                footer_blocks = [section.footer, section.first_page_footer, section.even_page_footer]

                for block in header_blocks + footer_blocks:
                    if not block:
                        continue
                    for para in block.paragraphs:
                        line = para.text.strip()
                        if line and line not in seen_lines:
                            # Keep header/footer details near top to help contact extraction.
                            sections_text.insert(0, line)
                            seen_lines.add(line)
        except Exception:
            pass

        # --- PASS 5: Extract hyperlink text/targets (handles mailto links and icon links) ---
        try:
            from docx.oxml.ns import qn

            rels = doc.part.rels

            def _add_link_target(target: str):
                if not target:
                    return
                normalized = target.strip()
                if normalized.lower().startswith('mailto:'):
                    normalized = normalized.split(':', 1)[1].split('?', 1)[0].strip()
                if normalized and normalized not in seen_lines:
                    # Keep contact links near top for downstream extraction.
                    sections_text.insert(0, normalized)
                    seen_lines.add(normalized)

            # Direct relationship scan catches links not represented as visible paragraph text.
            for rel in rels.values():
                target = getattr(rel, 'target_ref', '')
                if isinstance(target, str) and target:
                    if target.lower().startswith('mailto:'):
                        _add_link_target(target)

            # Paragraph hyperlink nodes: include display text + resolved target.
            for para in doc.paragraphs:
                for h in para._p.xpath('.//w:hyperlink'):
                    display_text = ''.join(t.text for t in h.xpath('.//w:t') if t.text).strip()
                    if display_text and display_text not in seen_lines:
                        sections_text.append(display_text)
                        seen_lines.add(display_text)

                    rid = h.get(qn('r:id'))
                    if rid and rid in rels:
                        _add_link_target(getattr(rels[rid], 'target_ref', ''))
        except Exception:
            pass

        if name_candidate and name_candidate not in sections_text[:5]:
            sections_text.insert(0, name_candidate)

        if not sections_text:
            return ""

        cleaned_lines = []
        for line in sections_text:
            line = re.sub(r'[^\w\s.,;:\'"@()\-#]', '', line).strip()
            line = re.sub(r'\s+', ' ', line).strip()
            if line:
                cleaned_lines.append(line)

        return "\n".join(cleaned_lines)
    
    except Exception as e:
        raise ValueError(f"Error parsing DOCX: {str(e)}")


def parse_txt(file):
    """
    Extract text from TXT file.
    Returns extracted text string.
    """
    try:
        content = file.read().decode('utf-8', errors='ignore')
        return clean_text(content) if content else ""
    
    except Exception as e:
        raise ValueError(f"Error parsing TXT: {str(e)}")


def parse_resume(file):
    """
    Parse resume file and extract text.
    Determines file type and calls appropriate parser.
    
    Args:
        file: werkzeug FileStorage object
    
    Returns:
        str: Extracted and cleaned resume text
    
    Raises:
        ValueError: If file type is not supported or parsing fails
    """
    filename = secure_filename(file.filename)
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    
    # Reset file pointer to beginning
    file.seek(0)
    
    if ext == 'pdf':
        return parse_pdf(file)
    elif ext == 'docx':
        return parse_docx(file)
    elif ext == 'txt':
        return parse_txt(file)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def extract_text(source):
    """
    Backward-compatible text extraction entrypoint.

    Accepts either:
    - a filesystem path to a resume file, or
    - a Werkzeug-like uploaded file object (with filename/read/seek).
    """
    if isinstance(source, (str, os.PathLike)):
        path = os.fspath(source)
        filename = os.path.basename(path)
        with open(path, "rb") as f:
            payload = f.read()

        class _MemoryFile(io.BytesIO):
            def __init__(self, data, name):
                super().__init__(data)
                self.filename = name

        return parse_resume(_MemoryFile(payload, filename))

    if hasattr(source, "filename") and hasattr(source, "read") and hasattr(source, "seek"):
        return parse_resume(source)

    raise ValueError("Unsupported input for extract_text; expected file path or uploaded file object")
